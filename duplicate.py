import os
import hashlib
from PyPDF2 import PdfReader
from docx import Document
from difflib import SequenceMatcher
from concurrent.futures import ThreadPoolExecutor
import re
import hashlib
from functools import lru_cache


# Helper function to calculate file hash
def calculate_hash(filepath):
    hasher = hashlib.md5()
    with open(filepath, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

# Find duplicates in specific file types
def find_duplicates(directory, file_extension):
    files = {}
    duplicates = []

    for root, _, filenames in os.walk(directory):
        for filename in filenames:
            if filename.lower().endswith(file_extension):
                filepath = os.path.join(root, filename)
                filehash = calculate_hash(filepath)

                if filehash in files:
                    duplicates.append((files[filehash], filepath))
                else:
                    files[filehash] = filepath

    return duplicates

# Find duplicates in any file type
def find_any_duplicates(directory):
    files = {}
    duplicates = []

    for root, _, filenames in os.walk(directory):
        for filename in filenames:
            filepath = os.path.join(root, filename)
            filehash = calculate_hash(filepath)

            if filehash in files:
                duplicates.append((files[filehash], filepath))
            else:
                files[filehash] = filepath

    return duplicates

# Function to process text content of various file types for similarity
def get_file_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif ext == '.docx':
        doc = Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


# Function to process text content of various file types for similarity
@lru_cache(maxsize=None)
def get_file_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        reader = PdfReader(filepath)
        return ''.join([page.extract_text() for page in reader.pages])
    elif ext == '.docx':
        doc = Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

def calculate_similarity(file_pair):
    file1, file2 = file_pair
    file1_text = get_file_text(file1)
    file2_text = get_file_text(file2)
    similarity = SequenceMatcher(None, file1_text, file2_text).ratio() * 100
    return {'file1': file1, 'file2': file2, 'similarity': similarity}

# Calculate similarity percentage among all files in the directory using parallel processing
def process_directory_similarity(directory):
    files = [os.path.join(root, filename)
             for root, _, filenames in os.walk(directory)
             for filename in filenames]

    similarities = []
    file_pairs = [(files[i], files[j]) for i in range(len(files)) for j in range(i + 1, len(files))]

    with ThreadPoolExecutor() as executor:
        results = list(executor.map(calculate_similarity, file_pairs))

    return results

